import re
import streamlit as st

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from pathlib import Path
from selenium import webdriver
from selenium.webdriver.chrome.options import Options


STRING_PATTERN_LINKS = r'IdLicitacao=(\d+)&IdEntidade=(\d+)&NrAnoLicitacao=(\d+)'
IMAGE_FOLDER = Path('images')

if not IMAGE_FOLDER.exists():
    IMAGE_FOLDER.mkdir()

if 'k_links' not in st.session_state:
    st.session_state['k_links'] = 2

@st.cache_data
def capture_screenshot(links):
    options = Options()
    options.add_argument("--start-maximized")

    # Configurar o caminho para o webdriver (certifique-se de ter o webdriver correspondente)
    # driver = webdriver.Chrome(executable_path=st.secrets.chromedriver_path, options=chrome_options)
    values = list()
    for link in links:
        url = link['link']
        save_path = IMAGE_FOLDER / f"{link['IdLicitacao']}.png"
        driver = webdriver.Chrome(options=options)
        try:
            driver.get(url)
            driver.implicitly_wait(2)
            # time.sleep(2)
            driver.save_screenshot(save_path)
            # window_height = driver.execute_script("return window.innerHeight;")
            values.append({
                'link' : url,
                'image' : str(save_path)
            })
        except Exception as e:
            print(f"Erro ao capturar screenshot de {url}: {e}")
        finally:
            driver.quit()

    print('- - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - ')
    return values

def ui_input_text_links(k, pattern=STRING_PATTERN_LINKS):

    def clear_links():
        st.toast('Ainda não implementado!', icon='😎')

    container = st.container()

    col1, col2, _ = st.columns([0.2, 0.2, 0.6])

    nro = col1.number_input(
        "Número de links para analisar",
        min_value=1,
        max_value=15,
        value=2,
        step=1,
        label_visibility="collapsed",
        )

    st.session_state['k_links'] = nro
    for i in range(nro):
        container.text_input("Link", key=f"link_{i}", label_visibility='collapsed')

    links = [st.session_state[f'link_{i}'] for i in range(k) ]
    matches = [ re.search(pattern, link) for link in  links ]

    with  col2:
        st.download_button(
            label="salvar links",
            data="\n".join(str(l) for l in links),
            file_name='links.txt',
            mime='txt/csv',
            disabled=not bool(len(links))
            )

    links = [{
        'seq' : i,
        'link' : link,
        'IdLicitacao': match.group(1),
        'IdEntidade': match.group(2),
        'NrAnoLicitacao': match.group(3),
    } for i, (link, match) in enumerate(zip(links, matches), start=1) if match is not None]

    return links

def save_to_docx(contents, save_path):

    doc = Document('template.docx')
    # section = doc.add_section()
    # section.orientation = WD_ORIENT.LANDSCAPE
    # section.left_margin = Inches(0.5)
    for content in contents:
        doc.add_page_break()
        p = doc.add_paragraph()
        doc.add_picture(content['image'], width=Inches(11.0))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_text(content['link'])
        font = run.font
        font.size = Pt(9)
        font.name = 'Arial'

        # print(section.orientation, section.page_width, section.page_height)
    doc.save(save_path)


def main():
    st.header("Screenshoter!")

    links = ui_input_text_links(st.session_state['k_links'])

    if st.button("Câmera, ação...", use_container_width=True) and len(links) > 0:
        values = capture_screenshot(links)
        docx = Path('relatorio.docx')
        if values is not None: save_to_docx(values, str(docx))
        if docx.exists():
            with open(docx, 'rb') as f:
                st.download_button("Download", f, file_name='relatorio.docx')



if __name__ == "__main__":
    main()